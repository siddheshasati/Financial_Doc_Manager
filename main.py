from fastapi import FastAPI, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
import models, auth, rag_service
from database import engine, get_db
from fastapi.security import OAuth2PasswordRequestForm
import pandas as pd
import io
import fitz
from docx import Document as DocxDocument
from sqlalchemy.orm import sessionmaker
from sqlalchemy.orm import Session
from enum import IntEnum

from passlib.context import CryptContext
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

models.Base.metadata.create_all(bind=engine)
app = FastAPI(title="Financial Doc Management")

SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)


# --- AUTH ENDPOINTS ---

@app.post("/auth/register")
async def register(
    username: str, 
    password: str,
    role_id: int,  
    db: Session = Depends(get_db)
):
    
    existing_user = db.query(models.User).filter(models.User.username == username).first()
    if existing_user:
        raise HTTPException(status_code=400, detail="Username already registered")

    
    role = db.query(models.Role).filter(models.Role.id == role_id).first()
    if not role:
        raise HTTPException(
            status_code=400, 
            detail="Invalid Role ID. Use: 1:Admin, 2:Analyst, 3:Auditor, 4:Client"
        )

    hashed_pwd = auth.get_password_hash(password) 

    new_user = models.User(
        username=username,
        password_hash=hashed_pwd,
        role_id=role_id  
    )
    
    try:
        db.add(new_user)
        db.commit()
        return {"message": f"User created successfully as {role.name}"}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail="Database error")

@app.post("/auth/login")
async def login(
    form_data: OAuth2PasswordRequestForm = Depends(), 
    db: Session = Depends(get_db)
):
    user = db.query(models.User).filter(models.User.username == form_data.username).first()
    
    
    if not user or not auth.verify_password(form_data.password, user.password_hash):
        raise HTTPException(status_code=400, detail="Incorrect username or password")
    
    access_token = auth.create_access_token(data={"sub": user.username})
    return {"access_token": access_token, "token_type": "bearer"}



@app.post("/documents/upload")
async def upload_document(
    title: str,
    doc_type: str,
    file: UploadFile = File(...),
    current_user: models.User = Depends(auth.get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role.name not in ["Admin", "Financial Analyst"]:
        raise HTTPException(status_code=403, detail="Not authorized to upload")
    
    content = await file.read()
    filename = file.filename.lower()  #type: ignore
    
    try:
        if filename.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(io.BytesIO(content))
            text_content = df.to_string()
        elif filename.endswith('.pdf'):
            with fitz.open(stream=content, filetype="pdf") as doc:
                text_content = "".join([page.get_text() for page in doc])  #type: ignore
        elif filename.endswith('.docx'):
            doc = DocxDocument(io.BytesIO(content))
            text_content = "\n".join([para.text for para in doc.paragraphs])
        else:
            text_content = content.decode("utf-8")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse {filename}: {str(e)}")

    new_doc = models.Document(
        title=title,
        document_type=doc_type,
        uploaded_by=current_user.username
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)
    
    rag_service.process_and_index(
        text_content,
        {"doc_id": new_doc.id, "title": title}
    )
    
    return {"message": f"{filename} indexed successfully", "doc_id": new_doc.id}


@app.post("/rag/search")
def search_insights(
    query: str, 
    current_user: models.User = Depends(auth.get_current_user)
):
    
    allowed_roles = ["Admin", "Auditor"]
    
    if not current_user.role or current_user.role.name not in allowed_roles:
        raise HTTPException(
            status_code=403, 
            detail="Search restricted to Admin or Auditor roles only"
        )
    
    
    try:
        results = rag_service.semantic_search(query)
        return {"results": results}
    except Exception as e:

        raise HTTPException(status_code=500, detail=f"Search service error: {str(e)}")

class RoleSelection(IntEnum):
    Admin = 1
    Financial_Analyst = 2
    Auditor = 3
    Client = 4

def seed_roles():
    db = SessionLocal()
    try:
        for role_choice in RoleSelection:
            
            role_exists = db.query(models.Role).filter(
                (models.Role.id == role_choice.value) | 
                (models.Role.name == role_choice.name)
            ).first()
            
            if not role_exists:
                new_role = models.Role(id=role_choice.value, name=role_choice.name)
                db.add(new_role)
                print(f"Seeded role: {role_choice.name}")
        
        db.commit()
    except Exception as e:
        print(f"Error seeding roles: {e}")
        db.rollback()
    finally:
        db.close()
